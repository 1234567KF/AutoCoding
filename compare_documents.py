from docx import Document

print("=== 对比两个文档 ===\n")

# 旧版本（python-docx 生成）
doc_old = Document(r'd:\projects\AI编程智驾\AI编程智驾手册.docx')
print(f"旧版本 (python-docx):")
print(f"  段落数: {len(doc_old.paragraphs)}")
print(f"  表格数: {len(doc_old.tables)}")

# 新版本（Pandoc + 模板）
doc_new = Document(r'd:\projects\AI编程智驾\AI编程智驾手册_模板版.docx')
print(f"\n新版本 (Pandoc + 模板):")
print(f"  段落数: {len(doc_new.paragraphs)}")
print(f"  表格数: {len(doc_new.tables)}")

# 样式统计
print("\n=== 新版本样式统计 ===\n")
styles_count = {}
for para in doc_new.paragraphs:
    style_name = para.style.name
    if style_name in styles_count:
        styles_count[style_name] += 1
    else:
        styles_count[style_name] = 1

for style, count in sorted(styles_count.items(), key=lambda x: x[1], reverse=True)[:10]:
    print(f"  {style}: {count}")

# 检查关键内容
print("\n=== 文档前10个段落 ===\n")
for i, para in enumerate(doc_new.paragraphs[:10]):
    if para.text.strip():
        text_preview = para.text[:80] + '...' if len(para.text) > 80 else para.text
        print(f"{i+1:2}. [{para.style.name:15}] {text_preview}")

# 检查表格
if len(doc_new.tables) > 0:
    print(f"\n=== 第一个表格预览 ===\n")
    table = doc_new.tables[0]
    print(f"表格大小: {len(table.rows)} 行 × {len(table.columns)} 列")
    if len(table.rows) > 0 and len(table.columns) > 0:
        print("第一行:")
        for cell in table.rows[0].cells:
            print(f"  | {cell.text[:30]}", end='')
        print(" |")
