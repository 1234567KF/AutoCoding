"""
Markdown 转 DOCX 转换器 - 参考模板样式
使用 markdown-to-docx 技能进行基础转换，然后应用模板样式
"""

import sys
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# 添加 markdown-to-docx 技能路径
skill_dir = r'C:\Users\KF\.trae-cn\skills\markdown-to-docx'
sys.path.insert(0, skill_dir)

from md2docx import MarkdownToDocxConverter

def convert_markdown_to_docx_with_template(
    input_md: str,
    output_docx: str,
    template_docx: str = None
):
    """
    将 Markdown 文件转换为 DOCX，并应用模板样式

    Args:
        input_md: 输入的 Markdown 文件路径
        output_docx: 输出的 DOCX 文件路径
        template_docx: 参考模板文件路径（可选）
    """

    # 第一步：使用 markdown-to-docx 技能进行基础转换
    print(f'正在使用 markdown-to-docx 技能转换 Markdown...')
    converter = MarkdownToDocxConverter(mermaid_theme='default')
    converter.convert_file(input_md, output_docx)
    print(f'✓ 基础转换完成: {output_docx}')

    # 第二步：应用模板样式
    if template_docx and os.path.exists(template_docx):
        print(f'\n正在应用模板样式...')
        apply_template_styles(output_docx, template_docx)
        print(f'✓ 模板样式应用完成')
    else:
        print(f'\n未找到模板文件，跳过样式应用')

    return output_docx

def apply_template_styles(output_docx: str, template_docx: str):
    """
    将模板样式应用到生成的 DOCX 文件

    Args:
        output_docx: 需要应用样式的 DOCX 文件
        template_docx: 模板 DOCX 文件
    """

    # 打开生成的文档
    doc = Document(output_docx)
    template = Document(template_docx)

    # 获取模板样式
    template_styles = {}
    for style_name in ['Heading 1', 'Heading 2', 'Normal']:
        if style_name in [s.name for s in template.styles]:
            template_styles[style_name] = template.styles[style_name]

    # 应用 Heading 1 样式到文档中的所有标题
    for para in doc.paragraphs:
        if para.style.name.startswith('Heading'):
            heading_level = para.style.name

            if heading_level == 'Heading 1' and 'Heading 1' in template_styles:
                apply_heading1_style(para, template_styles['Heading 1'])
            elif heading_level == 'Heading 2' and 'Heading 2' in template_styles:
                apply_heading2_style(para, template_styles['Heading 2'])
            elif heading_level == 'Heading 3' and 'Heading 2' in template_styles:
                # Heading 3 暂时使用 Heading 2 样式
                apply_heading2_style(para, template_styles['Heading 2'])

    # 应用 Normal 样式到正文段落
    for para in doc.paragraphs:
        if para.style.name == 'Normal':
            apply_normal_style(para, template_styles.get('Normal'))

    # 应用表格样式
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    # 设置表格文本样式
                    for run in para.runs:
                        run.font.name = '宋体'
                        run.font.size = Pt(9.375)
                        # 设置中文字体
                        r = run._element
                        r.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 保存文档
    doc.save(output_docx)

def apply_heading1_style(para, template_style):
    """应用 Heading 1 样式"""
    para.style.font.name = '黑体'
    para.style.font.size = Pt(15)
    para.style.font.bold = True
    para.style.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    para.style.paragraph_format.space_before = Pt(11)
    para.style.paragraph_format.space_after = Pt(11)

    # 设置中文字体
    para.style.font.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

def apply_heading2_style(para, template_style):
    """应用 Heading 2 样式"""
    para.style.font.name = '微软雅黑'
    para.style.font.size = Pt(12.5)
    para.style.paragraph_format.space_before = Pt(11.5)
    para.style.paragraph_format.space_after = Pt(11.5)

def apply_normal_style(para, template_style):
    """应用 Normal 样式到正文"""
    para.style.font.name = 'Times New Roman'
    para.style.font.size = Pt(10.75)
    # 首行缩进 2 个字符（约 46.5pt）
    para.style.paragraph_format.first_line_indent = Emu(662940)
    para.style.paragraph_format.line_spacing = 1.5

    # 设置中文字体
    para.style.font.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

if __name__ == '__main__':
    # 设置路径
    input_md = r'd:\projects\AI编程智驾\AI编程智驾手册.md'
    output_docx = r'd:\projects\AI编程智驾\AI编程智驾手册.docx'
    template_docx = r'd:\projects\AI编程智驾\标准文件.docx'

    print('=' * 60)
    print('Markdown 转 DOCX 转换工具')
    print('参考模板样式进行转换')
    print('=' * 60)
    print()

    # 执行转换
    result = convert_markdown_to_docx_with_template(
        input_md=input_md,
        output_docx=output_docx,
        template_docx=template_docx
    )

    print()
    print('=' * 60)
    print(f'✓ 转换完成！')
    print(f'输出文件: {result}')
    print('=' * 60)
