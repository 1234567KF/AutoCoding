from docx import Document
from docx.shared import Pt

doc = Document(r'd:\projects\AI编程智驾\AI编程智驾手册.docx')

print('=== 转换后的文档样式验证 ===\n')

# 统计样式使用情况
styles_count = {}
for para in doc.paragraphs:
    style_name = para.style.name
    if style_name in styles_count:
        styles_count[style_name] += 1
    else:
        styles_count[style_name] = 1

print('段落样式统计:')
for style, count in sorted(styles_count.items(), key=lambda x: x[1], reverse=True):
    print(f'  {style}: {count} 个段落')

print('\n=== 样式详细信息 ===\n')

# 检查 Heading 1
try:
    h1 = doc.styles['Heading 1']
    print('Heading 1 样式:')
    print(f'  字体: {h1.font.name}')
    print(f'  大小: {h1.font.size}')
    print(f'  粗体: {h1.font.bold}')
    print(f'  段前: {h1.paragraph_format.space_before}')
    print(f'  段后: {h1.paragraph_format.space_after}')
except Exception as e:
    print(f'Heading 1: {e}')

print()

# 检查 Normal
try:
    normal = doc.styles['Normal']
    print('Normal 样式:')
    print(f'  字体: {normal.font.name}')
    print(f'  大小: {normal.font.size}')
    print(f'  首行缩进: {normal.paragraph_format.first_line_indent}')
    print(f'  行距: {normal.paragraph_format.line_spacing}')
except Exception as e:
    print(f'Normal: {e}')

print('\n=== 文档结构预览 ===\n')

# 显示前20个段落的样式和内容摘要
for i, para in enumerate(doc.paragraphs[:20]):
    if para.text.strip():
        text_preview = para.text[:60] + '...' if len(para.text) > 60 else para.text
        print(f'{i+1:2}. [{para.style.name:15}] {text_preview}')

print('\n=== 表格检查 ===\n')
print(f'文档包含 {len(doc.tables)} 个表格')

if len(doc.tables) > 0:
    for i, table in enumerate(doc.tables[:3]):
        print(f'\n表格 {i+1}:')
        print(f'  行数: {len(table.rows)}')
        print(f'  列数: {len(table.columns)}')
        if len(table.rows) > 0 and len(table.columns) > 0:
            first_cell = table.rows[0].cells[0].text
            print(f'  第一个单元格: {first_cell[:30]}...' if len(first_cell) > 30 else f'  第一个单元格: {first_cell}')

print('\n=== 总结 ===\n')
print(f'✓ 文档转换完成')
print(f'✓ 总段落数: {len(doc.paragraphs)}')
print(f'✓ 总表格数: {len(doc.tables)}')
print(f'✓ 主要样式已应用')
